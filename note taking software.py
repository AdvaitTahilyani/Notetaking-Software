#note taking software 
import os
import whisper
import time as time
import subprocess 
from pydub import AudioSegment
import shutil
from flask import Flask, jsonify, request
import csv
from langchain_community.llms import Ollama
from docx import Document 
import ssl
import urllib.request
import whisper

ssl._create_default_https_context = ssl._create_unverified_context


def create_word_doc(content):
    first_sentence_end = content.find('\n')
    if first_sentence_end != -1:
        title = content[:first_sentence_end + 1]  # +1 to include the period
    else:
        title = "Summary"

    doc = Document()  # Use Document() from python-docx
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    filename = f"{title.removesuffix('.').removeprefix('Title: ')}.docx"
    doc.save(filename)
    return filename

model = whisper.load_model("small.en")

# app = Flask(__name__)

# @app.route('/process', methods=['POST'])
# def process():
#     data = request.json
#     # Your backend logic here
#     result = {"status": "success", "data": data}
#     return jsonify(result)

# if __name__ == '__main__':
#     app.run(port=5000)

model = whisper.load_model("small.en")
# recordings = "Library/Application Support/com.apple.voicememos/Recordings"
folder = "Recordings"
destination = "Transcripts"
# backup = "Backup"
# recs = os.listdir(recordings)
# for i in recs:
#     if (".m4a" in i ):
#         shutil.move(os.path.join(recordings, i), os.path.join(folder, i))
#         shutil.copyfile(os.path.join(recordings, i), os.path.join(backup, i))
# files = os.listdir(folder)
# files = [os.path.join(folder, f) for f in files] # add path to each file
# for f in files:
#     os.rename(f,f.replace("m4a","wav"))
files = os.listdir(folder)
files = [os.path.join(folder, f) for f in files] # add path to each filenformation is
# for i in files:
#     t = 10 * 60 * 100
#     if ".wav" not in i:
#         continue
#     t1 = 0 #Works in milliseconds
#     t2 = t
#     print(i)
#     try:
#         newAudio = AudioSegment.from_wav(i)
#     except:
#         newAudio = AudioSegment.from_file(i, format="mp4")
#     print(len(newAudio))
#     for j in range(int(len(newAudio)/t)):
#         newAudio1 = newAudio[t1:t2]
#         newAudio1.export(i.replace(".wav", str(j)) + ".wav", format="wav")
#         t1 += t
#         t2 += t

files.sort(key=lambda x: os.path.getmtime(x), reverse=True)
for i in files:
    if not ".wav" in i:
        continue
    result = model.transcribe(i)
    with open(os.path.join(destination, i[len(folder)+1:(len(i)-3)]+"txt"), "w") as f:
        f.write(result["text"])
    # os.remove(i)

    LLM = Ollama(model="phi3")
    prompt = result["text"] + "Please summarize the following in a manner that the suppositions and meaning of the text are maintained and every single important detail is retained. Please take the liberty to add in any information you believe will improve the understanding of the summary and add titles and sub titles so that the summary is organized"
    output = LLM.invoke(prompt)
    doc_filename = create_word_doc(output)
    print(f"Document created with filename: {doc_filename}")
