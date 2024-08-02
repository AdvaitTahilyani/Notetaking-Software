import os
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import whisper
from docx import Document
import ssl
from langchain_community.llms import Ollama

ssl._create_default_https_context = ssl._create_unverified_context

app = Flask(__name__)
model = whisper.load_model("small.en")

UPLOAD_FOLDER = 'uploads'  # Keep uploads separate
DOWNLOADS_FOLDER = os.path.expanduser('~/Downloads')  # Use the user's Downloads directory
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOADS_FOLDER, exist_ok=True)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        transcript_path, docx_path = process_audio(file_path)
        return jsonify({
            'message': 'Transcription completed',
            'transcript_file_path': transcript_path,
            'docx_file_path': docx_path
        })

def process_audio(file_path):
    base_filename = os.path.basename(file_path)
    name_only = os.path.splitext(base_filename)[0]
    result = model.transcribe(file_path)
    text_filename = name_only + '.txt'
    text_file_path = os.path.join(DOWNLOADS_FOLDER, text_filename)
    with open(text_file_path, 'w') as f:
        f.write(result['text'])

    LLM = Ollama(model="phi3")
    prompt = result["text"] + "Please summarize the following in a manner that the suppositions and meaning of the text are maintained and every single important detail is retained. Please take the liberty to add in any information you believe will improve the understanding of the summary and add titles and sub titles so that the summary is organized"
    output = LLM.invoke(prompt)
    doc_filename = create_word_doc(output, name_only)
    doc_file_path = os.path.join(DOWNLOADS_FOLDER, doc_filename)
    return text_file_path, doc_file_path

def create_word_doc(content, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    filename = f"{title}.docx"
    doc_path = os.path.join(DOWNLOADS_FOLDER, filename)
    doc.save(doc_path)
    return doc_path

if __name__ == '__main__':
    app.run(debug=True, port=5000)
